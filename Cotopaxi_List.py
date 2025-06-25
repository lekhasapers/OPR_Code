import pandas as pd
import datetime as dt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import gspread
from oauth2client.service_account import ServiceAccountCredentials

# 1) read only dw:-)
scope = [
    'https://www.googleapis.com/auth/spreadsheets.readonly',
    'https://www.googleapis.com/auth/drive.metadata.readonly',
]
creds  = ServiceAccountCredentials.from_json_keyfile_name('credentials.json', scope)
client = gspread.authorize(creds)

sheet = client.open('Cotopaxi Media Report - 2025').worksheet('ONLINE')
rows  = sheet.get_all_values()[1:]       # skip header
df    = pd.DataFrame(rows)

def parse_impr(r):
    for i in (5, 6):
        v = r[i].replace(',', '').strip()
        if v.isdigit():
            return int(v)
    return 0

df['Impressions'] = df.apply(parse_impr, axis=1)

# cat logic
def get_category(r):
    if r[8]=='P' and r[9]=='A': return 'APPAREL FEATURES'
    if r[8]=='P' and r[9]=='P': return 'PACKS & BAG FEATURES'
    if r[8]=='B':               return 'BRAND/SUSTAINABILITY'
    return None

df['Category'] = df.apply(get_category, axis=1)

today        = dt.date.today()
delta_to_wed = (2 - today.weekday()) % 7  # Wed=2
end_date     = today + dt.timedelta(days=delta_to_wed)
start_date   = end_date - dt.timedelta(days=6)

df[4] = pd.to_datetime(df[4], errors='coerce')
week_df = df[df[4].dt.date.between(start_date, end_date)]
week_df = week_df.sort_values(['Category', 4], ascending=[True, True])

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    rid  = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), rid)
    r     = OxmlElement('w:r')
    rPr   = OxmlElement('w:rPr')
    c     = OxmlElement('w:color');    c.set(qn('w:val'), "0000FF"); rPr.append(c)
    u     = OxmlElement('w:u');        u.set(qn('w:val'), "single");  rPr.append(u)
    r.append(rPr)
    t     = OxmlElement('w:t');        t.text = text
    r.append(t)
    link.append(r)
    paragraph._p.append(link)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

hdr = doc.add_paragraph()
hdr.add_run(f"Date Range: {start_date:%B %d, %Y} – {end_date:%B %d, %Y}").bold = True
doc.add_paragraph(f"Total Impressions: {week_df['Impressions'].sum():,}")
doc.add_paragraph(f"Total Hits: {len(week_df)}")
doc.add_paragraph("")  

for section in ['APPAREL FEATURES','PACKS & BAG FEATURES','BRAND/SUSTAINABILITY']:
    subset = week_df[week_df['Category']==section]
    if subset.empty:
        continue

    sec_p = doc.add_paragraph()
    sec_p.add_run(section).bold = True

    for _, row in subset.iterrows():
        outlet_domain = row[0]
        base = outlet_domain.split('/')[0].split('.')[-2]
        outlet = ' '.join(base.replace('-', ' ').split()).title()

        title = row[3]
        url   = row[7]
        imp   = row['Impressions']

        p = doc.add_paragraph()
        run = p.add_run(f"{outlet}: ")
        run.bold = True
        add_hyperlink(p, url, title)
        p.add_run(f" - {imp:,} Impressions")

    # extra space between categories
    doc.add_paragraph("")

# output!
output = 'cotopaxi_weekly_hits.docx'
doc.save(output)
print("Generated:", output)
